"""
backend/export_generator.py

ExportGenerator: generates downloadable exports in 8 formats.

Export types:
  graph-json        — Full graph as JSON
  graph-csv         — Nodes + edges as ZIP of two CSV files
  bibtex            — BibTeX citations for all papers
  literature-review — LLM-generated Markdown literature review
  genealogy-pdf     — Genealogy narrative as formatted PDF (WeasyPrint)
  action-log        — Session action history as JSON
  graph-png         — Static graph image (matplotlib, 150dpi)
  graph-svg         — Graph as SVG

All generate() calls:
  1. Build the export in memory
  2. Upload to R2 at exports/{session_id}/{timestamp}_{filename}
  3. Return a presigned URL (1 hour)

Note on nodes_by_id: export_to_json() returns nodes as a list.
Always use _build_nodes_by_id(graph_data) — never access graph_data['nodes_by_id'].
"""
import csv
import io
import json
import logging
import zipfile
from datetime import datetime, timezone
from typing import Optional

from backend.r2_client import R2Client
from backend.db import fetchall

logger = logging.getLogger(__name__)


def _build_nodes_by_id(graph_data: dict) -> dict:
    """Build paper_id → node dict from the nodes list."""
    return {node["id"]: node for node in graph_data.get("nodes", [])}


class ExportGenerator:
    """Stateless — instantiate fresh per request."""

    EXPORT_TYPES = [
        "graph-json", "graph-csv", "bibtex", "literature-review",
        "genealogy-pdf", "action-log", "graph-png", "graph-svg",
    ]

    def __init__(self):
        self.r2 = R2Client()

    def generate(
        self,
        export_type:  str,
        graph_data:   dict,
        session_id:   str,
        llm_client=None,
        extra:        Optional[dict] = None,
    ) -> str:
        """
        Generate an export, upload to R2, and return a presigned download URL.
        Returns presigned URL string (valid 1 hour).
        Raises ValueError for unknown export_type.
        Raises RuntimeError if R2 is not configured.
        """
        if export_type not in self.EXPORT_TYPES:
            raise ValueError(
                f"Unknown export type: {export_type!r}. Allowed: {self.EXPORT_TYPES}"
            )
        if not self.r2._enabled:
            raise RuntimeError(
                "R2 storage is not configured — cannot generate downloadable exports. "
                "Set R2_ACCOUNT_ID, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY in environment."
            )

        extra = extra or {}
        generators = {
            "graph-json":        self._graph_json,
            "graph-csv":         self._graph_csv,
            "bibtex":            self._bibtex,
            "literature-review": self._literature_review,
            "genealogy-pdf":     self._genealogy_pdf,
            "action-log":        self._action_log,
            "graph-png":         self._graph_png,
            "graph-svg":         self._graph_svg,
        }

        file_bytes, filename, content_type = generators[export_type](
            graph_data, session_id, llm_client, extra
        )

        ts  = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        key = f"exports/{session_id}/{ts}_{filename}"
        self.r2.upload(key, file_bytes, content_type)

        return self.r2.presigned_url(key, expires_in=3600)

    # ── Format generators ─────────────────────────────────────────────────────

    def _graph_json(self, graph_data, session_id, llm_client, extra):
        content = json.dumps(graph_data, indent=2, ensure_ascii=False).encode("utf-8")
        return content, "arivu_graph.json", "application/json"

    def _graph_csv(self, graph_data, session_id, llm_client, extra):
        nodes_by_id = _build_nodes_by_id(graph_data)
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            nodes_buf  = io.StringIO()
            fieldnames = ["paper_id","title","authors","year","citation_count",
                          "fields_of_study","is_seed","url","doi"]
            writer = csv.DictWriter(nodes_buf, fieldnames=fieldnames, extrasaction="ignore")
            writer.writeheader()
            for node in graph_data.get("nodes", []):
                writer.writerow({
                    "paper_id":        node.get("id", ""),
                    "title":           node.get("title", ""),
                    "authors":         "; ".join(node.get("authors", [])),
                    "year":            node.get("year", ""),
                    "citation_count":  node.get("citation_count", ""),
                    "fields_of_study": "; ".join(node.get("fields_of_study", [])),
                    "is_seed":         node.get("is_seed", False),
                    "url":             node.get("url", ""),
                    "doi":             node.get("doi", ""),
                })
            zf.writestr("nodes.csv", nodes_buf.getvalue())

            edges_buf   = io.StringIO()
            edge_fields = ["citing_paper_id","cited_paper_id","citing_title","cited_title",
                           "mutation_type","confidence_tier","similarity_score",
                           "citing_sentence","cited_sentence"]
            writer = csv.DictWriter(edges_buf, fieldnames=edge_fields, extrasaction="ignore")
            writer.writeheader()
            for edge in graph_data.get("edges", []):
                src_id = edge.get("source") or edge.get("citing_paper_id", "")
                tgt_id = edge.get("target") or edge.get("cited_paper_id", "")
                writer.writerow({
                    "citing_paper_id":  src_id,
                    "cited_paper_id":   tgt_id,
                    "citing_title":     nodes_by_id.get(src_id, {}).get("title", ""),
                    "cited_title":      nodes_by_id.get(tgt_id, {}).get("title", ""),
                    "mutation_type":    edge.get("mutation_type", ""),
                    "confidence_tier":  edge.get("confidence_tier", ""),
                    "similarity_score": edge.get("similarity_score", ""),
                    "citing_sentence":  edge.get("citing_sentence", ""),
                    "cited_sentence":   edge.get("cited_sentence", ""),
                })
            zf.writestr("edges.csv", edges_buf.getvalue())
        return buf.getvalue(), "arivu_graph.zip", "application/zip"

    def _bibtex(self, graph_data, session_id, llm_client, extra):
        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        meta  = graph_data.get("metadata", {})
        lines = [
            f"% Arivu citation export — {today}",
            f"% Seed paper: {meta.get('seed_paper_id', 'unknown')}",
            f"% Total papers: {len(graph_data.get('nodes', []))}",
            "",
        ]
        seen_keys: set = set()
        for node in graph_data.get("nodes", []):
            authors    = node.get("authors", []) or ["Unknown"]
            first_auth = (authors[0].split(",")[0].split()[-1] if authors[0] else "Unknown")
            year       = node.get("year", "nd") or "nd"
            base_key   = f"{first_auth}{year}"
            cite_key   = base_key
            suffix     = 1
            while cite_key in seen_keys:
                cite_key = f"{base_key}{chr(ord('a') + suffix - 1)}"
                suffix  += 1
            seen_keys.add(cite_key)

            title       = (node.get("title", "Untitled") or "Untitled").replace("{", r"\{").replace("}", r"\}")
            authors_str = " and ".join(authors)
            entry = [f"@article{{{cite_key},"]
            entry.append(f"  title  = {{{title}}},")
            entry.append(f"  author = {{{authors_str}}},")
            entry.append(f"  year   = {{{year}}},")
            if node.get("doi"):
                entry.append(f"  doi    = {{{node['doi']}}},")
            if node.get("url"):
                entry.append(f"  url    = {{{node['url']}}},")
            entry.append("}")
            lines.extend(entry)
            lines.append("")
        return "\n".join(lines).encode("utf-8"), "arivu_citations.bib", "text/plain"

    def _literature_review(self, graph_data, session_id, llm_client, extra):
        if llm_client and getattr(llm_client, "available", False):
            try:
                review_text = llm_client.generate_literature_review(graph_data)
            except Exception as exc:
                logger.warning(f"LLM literature review failed: {exc} — using template")
                review_text = self._template_literature_review(graph_data)
        else:
            review_text = self._template_literature_review(graph_data)
        return review_text.encode("utf-8"), "arivu_literature_review.md", "text/markdown"

    def _template_literature_review(self, graph_data: dict) -> str:
        meta       = graph_data.get("metadata", {})
        nodes      = graph_data.get("nodes", [])
        seed_title = meta.get("seed_paper_title", "the seed paper")
        today      = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        sorted_nodes = sorted(nodes, key=lambda n: n.get("year") or 9999)
        lines = [
            f"# Literature Review: {seed_title}",
            f"*Generated by Arivu — {today}*", "",
            "## Overview",
            f"This graph traces the intellectual ancestry of **{seed_title}**, "
            f"encompassing {len(nodes)} papers across "
            f"{len(set(f for n in nodes for f in n.get('fields_of_study', [])))} fields.",
            "", "## Foundational Papers", "",
        ]
        for node in sorted_nodes[:10]:
            authors    = (node.get("authors") or ["Unknown"])
            author_str = f"{authors[0]} et al." if len(authors) > 1 else authors[0]
            lines.append(f"- **{node['title']}** ({author_str}, {node.get('year', 'n.d.')})")
        lines.extend(["", "## Chronological Summary", ""])
        for node in sorted_nodes:
            authors    = (node.get("authors") or ["Unknown"])
            author_str = f"{authors[0]} et al." if len(authors) > 1 else authors[0]
            lines.append(f"- **{node.get('year', 'n.d.')}** — {node['title']} ({author_str})")
        return "\n".join(lines)

    def _genealogy_pdf(self, graph_data, session_id, llm_client, extra):
        if llm_client and getattr(llm_client, "available", False):
            try:
                result    = llm_client.generate_genealogy_story(graph_data)
                narrative = result.get("narrative", "") if isinstance(result, dict) else str(result)
            except Exception as exc:
                logger.warning(f"LLM genealogy failed: {exc} — using template")
                narrative = self._template_genealogy(graph_data)
        else:
            narrative = self._template_genealogy(graph_data)

        try:
            import markdown as md_lib
            from weasyprint import HTML as WeasyHTML
            html_body  = md_lib.markdown(narrative)
            meta       = graph_data.get("metadata", {})
            seed_title = meta.get("seed_paper_title", "Research Paper")
            today      = datetime.now(timezone.utc).strftime("%Y-%m-%d")
            html_full  = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: Georgia, serif; max-width: 680px; margin: 60px auto;
          font-size: 14px; line-height: 1.75; color: #1a1a2e; }}
  h1   {{ font-size: 22px; color: #0a0e17; margin-bottom: 0.3em; }}
  h2   {{ font-size: 16px; color: #334155; margin-top: 2em; }}
  p    {{ margin: 1em 0; }}
  .meta   {{ font-size: 12px; color: #64748B; margin-bottom: 2em; }}
  .footer {{ margin-top: 60px; padding-top: 16px; border-top: 1px solid #e2e8f0;
             font-size: 11px; color: #94A3B8; }}
</style></head><body>
<h1>{seed_title} — Intellectual Genealogy</h1>
<div class="meta">Generated by Arivu · {today}</div>
{html_body}
<div class="footer">
  Generated by Arivu — arivu.app · {len(graph_data.get('nodes', []))} papers analysed
</div>
</body></html>"""
            pdf_bytes = WeasyHTML(string=html_full).write_pdf()
        except ImportError:
            pdf_bytes = self._narrative_to_reportlab_pdf(narrative, graph_data)
        return pdf_bytes, "arivu_genealogy.pdf", "application/pdf"

    def _template_genealogy(self, graph_data: dict) -> str:
        meta   = graph_data.get("metadata", {})
        nodes  = graph_data.get("nodes", [])
        edges  = graph_data.get("edges", [])
        seed   = meta.get("seed_paper_title", "the seed paper")
        years  = [n.get("year") for n in nodes if n.get("year")]
        yr_range = f"{min(years)}–{max(years)}" if years else "unknown"
        return (
            f"# The Intellectual Ancestry of {seed}\n\n"
            f"This graph traces {len(nodes)} papers published between {yr_range}, "
            f"connected by {len(edges)} inheritance relationships.\n\n"
            f"*Full narrative requires Groq API key. "
            f"Set GROQ_API_KEY in environment to enable LLM-generated analysis.*"
        )

    def _narrative_to_reportlab_pdf(self, text: str, graph_data: dict) -> bytes:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        buf    = io.BytesIO()
        doc    = SimpleDocTemplate(buf, pagesize=letter,
                                   leftMargin=inch, rightMargin=inch,
                                   topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        story  = []
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.15 * inch))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], styles["Title"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            else:
                story.append(Paragraph(line, styles["Normal"]))
        doc.build(story)
        return buf.getvalue()

    def _action_log(self, graph_data, session_id, llm_client, extra):
        try:
            actions = fetchall(
                "SELECT action_type, action_data, timestamp FROM action_log "
                "WHERE session_id = %s ORDER BY timestamp DESC",
                (session_id,),
            )
        except Exception as exc:
            logger.warning(f"action_log fetch failed: {exc}")
            actions = []
        content = json.dumps({
            "session_id":   session_id,
            "exported_at":  datetime.now(timezone.utc).isoformat(),
            "action_count": len(actions),
            "actions": [
                {"action_type": a["action_type"], "action_data": a["action_data"],
                 "timestamp": str(a["timestamp"])}
                for a in actions
            ],
        }, indent=2, default=str).encode("utf-8")
        return content, "arivu_action_log.json", "application/json"

    def _graph_png(self, graph_data, session_id, llm_client, extra):
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import networkx as nx
        import math as _math

        nodes_by_id = _build_nodes_by_id(graph_data)
        G = nx.DiGraph()
        for node in graph_data.get("nodes", []):
            G.add_node(node["id"])
        for edge in graph_data.get("edges", []):
            src = edge.get("source") or edge.get("citing_paper_id", "")
            tgt = edge.get("target") or edge.get("cited_paper_id", "")
            if src and tgt:
                G.add_edge(src, tgt)

        fig = plt.figure(figsize=(16, 12), facecolor="#0a0e17")
        ax  = fig.add_subplot(111)
        ax.set_facecolor("#0a0e17")
        pos = nx.spring_layout(G, k=2.0, seed=42)
        node_sizes = [
            max(20, _math.log1p(nodes_by_id.get(n, {}).get("citation_count", 1) or 1) * 30)
            for n in G.nodes()
        ]
        nx.draw_networkx(G, pos=pos, ax=ax, with_labels=False,
                         node_size=node_sizes, node_color="#3B82F6",
                         edge_color="#475569", arrows=True, arrowsize=6, alpha=0.85)
        meta = graph_data.get("metadata", {})
        ax.set_title(f"Citation Ancestry: {meta.get('seed_paper_title', 'Graph')}",
                     color="#E2E8F0", fontsize=13, pad=12)
        ax.axis("off")
        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor="#0a0e17", edgecolor="none")
        plt.close(fig)
        return buf.getvalue(), "arivu_graph.png", "image/png"

    def _graph_svg(self, graph_data, session_id, llm_client, extra):
        if extra.get("svg_data"):
            svg_str = extra["svg_data"]
            if isinstance(svg_str, str):
                return svg_str.encode("utf-8"), "arivu_graph.svg", "image/svg+xml"

        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import networkx as nx

        G = nx.DiGraph()
        for node in graph_data.get("nodes", []):
            G.add_node(node["id"])
        for edge in graph_data.get("edges", []):
            src = edge.get("source") or edge.get("citing_paper_id", "")
            tgt = edge.get("target") or edge.get("cited_paper_id", "")
            if src and tgt:
                G.add_edge(src, tgt)

        fig = plt.figure(figsize=(16, 12), facecolor="#0a0e17")
        ax  = fig.add_subplot(111)
        ax.set_facecolor("#0a0e17")
        pos = nx.spring_layout(G, k=2.0, seed=42)
        nx.draw_networkx(G, pos=pos, ax=ax, with_labels=False,
                         node_color="#3B82F6", edge_color="#475569", alpha=0.85)
        ax.axis("off")
        buf = io.StringIO()
        plt.savefig(buf, format="svg", bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue().encode("utf-8"), "arivu_graph.svg", "image/svg+xml"

    # ── Phase 8: DOCX generation ───────────────────────────────────────────────

    def generate_docx(self, sections: list[dict], title: str = "Arivu Report") -> bytes:
        """
        Generate a Word document from structured sections.
        Each section: {"heading": str, "body": str}
        Returns raw .docx bytes.
        Used by LiteratureReviewEngine and FieldEntryKit.
        """
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        doc.add_heading(title, level=0)

        for section in sections:
            if section.get("heading"):
                doc.add_heading(section["heading"], level=1)
            if section.get("body"):
                for para_text in section["body"].split("\n\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
